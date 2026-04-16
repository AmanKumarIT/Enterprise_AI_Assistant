"""
Ingestion pipeline for file uploads (PDF, DOCX, TXT).
Parses documents, cleans text, and passes to the base pipeline
for chunking/embedding/storage.
"""
import io
import os
import logging
from typing import List

from sqlalchemy.future import select

from app.models.document import Document
from app.ingestion.base import BaseIngestionPipeline, RawDocument
from app.services.chunking import TextChunker

logger = logging.getLogger(__name__)


class FileIngestionPipeline(BaseIngestionPipeline):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def get_chunker(self):
        return TextChunker(chunk_size=1000, chunk_overlap=200)

    async def extract_documents(self) -> List[RawDocument]:
        raw_docs: List[RawDocument] = []

        # Query all active documents for this data source
        stmt = select(Document).where(
            Document.data_source_id == self.data_source.id,
            Document.is_active == True,
        )
        result = await self.db.execute(stmt)
        documents = result.scalars().all()
        logger.info("Fetched %d Document records for data_source_id=%s", len(documents), self.data_source.id)

        for doc in documents:
            if not doc.source_uri or not os.path.exists(doc.source_uri):
                logger.warning(
                    "Skipping document %s: file not found at %s",
                    doc.title, doc.source_uri
                )
                continue

            logger.info("Reading file: %s (path=%s)", doc.title, doc.source_uri)

            try:
                with open(doc.source_uri, "rb") as f:
                    content_bytes = f.read()

                text = self._parse_file(doc.title, content_bytes)
                if text and text.strip():
                    text = self._clean_text(text)
                    logger.info("Extracted %d characters of text from %s", len(text), doc.title)
                    raw_docs.append(
                        RawDocument(
                            title=doc.title,
                            content=text,
                            source_uri=doc.source_uri,
                            metadata={"original_filename": doc.title},
                            mime_type=doc.mime_type or self._get_mime_type(doc.title),
                        )
                    )
                else:
                    logger.warning("No text extracted from %s", doc.title)
            except Exception as e:
                logger.exception("Failed to parse file '%s'", doc.title)

        logger.info("Successfully produced %d RawDocument(s) for the pipeline", len(raw_docs))

        return raw_docs

    def _parse_file(self, filename: str, content: bytes) -> str:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return self._parse_pdf(content)
        elif lower.endswith(".docx"):
            return self._parse_docx(content)
        elif lower.endswith(".txt") or lower.endswith(".md") or lower.endswith(".rst"):
            return content.decode("utf-8", errors="replace")
        else:
            return content.decode("utf-8", errors="replace")

    def _parse_pdf(self, content: bytes) -> str:
        import pdfplumber

        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)

    def _parse_docx(self, content: bytes) -> str:
        import docx

        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _clean_text(self, text: str) -> str:
        import re

        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        text = re.sub(r"\t+", " ", text)
        return text.strip()

    def _get_mime_type(self, filename: str) -> str:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return "application/pdf"
        elif lower.endswith(".docx"):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif lower.endswith(".txt"):
            return "text/plain"
        elif lower.endswith(".md"):
            return "text/markdown"
        return "application/octet-stream"
